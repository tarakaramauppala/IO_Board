# -*- coding: utf-8 -*-
"""Build the editable Word (.docx) version of the VX-0057 Product & Safety Guide.

Shares all content with the HTML/PDF via content.py, so the two stay in sync.
Uses python-docx with shaded tables for the ANSI Z535 hazard panels and the
PNG icons rasterized from the same SVGs used in the PDF.
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import content as C

HERE = os.path.dirname(os.path.abspath(__file__))
ICON = os.path.join(HERE, "assets", "icons")
BOARD = os.path.join(HERE, "assets", "board.png")
OUT = os.path.join(HERE, "VX-0057-Product-Safety-Guide.docx")
M = C.META

NAVY = RGBColor(0x0E, 0x2A, 0x47)
TEAL = RGBColor(0x0F, 0x6F, 0xA3)
GREY = RGBColor(0x5B, 0x66, 0x75)
INK = RGBColor(0x1B, 0x24, 0x30)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BLACK = RGBColor(0x00, 0x00, 0x00)
LINE = "D9DEE6"
SOFT = "F4F6F9"


def ic(name):
    return os.path.join(ICON, f"{name}.png")


# ----------------------------------------------------------------------- xml
def _set_shade(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill_hex)
    tcPr.append(shd)


def _set_borders(el, color_hex, sz=4, sides=("top", "left", "bottom", "right")):
    """Borders on a table (el = table) — sz in eighths of a point."""
    tblPr = el._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for side in sides:
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color_hex)
        borders.append(e)
    # also kill inside borders
    for side in ("insideH", "insideV"):
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:val"), "nil")
        borders.append(e)
    tblPr.append(borders)


def _no_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{side}")
        e.set(qn("w:val"), "nil")
        borders.append(e)
    tblPr.append(borders)


def _cell_margins(cell, top=40, bottom=40, left=80, right=80):
    tcPr = cell._tc.get_or_add_tcPr()
    m = OxmlElement("w:tcMar")
    for name, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        e = OxmlElement(f"w:{name}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        m.append(e)
    tcPr.append(m)


def _cant_split(row):
    trPr = row._tr.get_or_add_trPr()
    cs = OxmlElement("w:cantSplit")
    cs.set(qn("w:val"), "true")
    trPr.append(cs)


def _keep_with_next(p):
    pPr = p._p.get_or_add_pPr()
    kn = OxmlElement("w:keepNext")
    pPr.append(kn)


def _para_bottom_border(p, color_hex=LINE, sz=10):
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    b = OxmlElement("w:bottom")
    b.set(qn("w:val"), "single")
    b.set(qn("w:sz"), str(sz))
    b.set(qn("w:space"), "4")
    b.set(qn("w:color"), color_hex)
    pbdr.append(b)
    pPr.append(pbdr)


def _field(paragraph, instr):
    run = paragraph.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
    run._r.append(f1); run._r.append(it); run._r.append(f2)
    return run


# ------------------------------------------------------------------- helpers
def run(p, text, size=10.5, bold=False, italic=False, color=INK, spacing=None, caps=False):
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    r.font.name = "Segoe UI"
    if caps:
        r.font.all_caps = True
    if spacing is not None:
        rPr = r._element.get_or_add_rPr()
        sp = OxmlElement("w:spacing"); sp.set(qn("w:val"), str(spacing))
        rPr.append(sp)
    return r


def heading(doc, text, size=13, rule=True, space_before=10, space_after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run(p, text, size=size, bold=True, color=NAVY)
    if rule:
        _para_bottom_border(p)
    _keep_with_next(p)
    return p


def set_row_height(row, inches):
    trPr = row._tr.get_or_add_trPr()
    h = OxmlElement("w:trHeight")
    h.set(qn("w:val"), str(int(inches * 1440)))
    h.set(qn("w:hRule"), "atLeast")
    trPr.append(h)


# --------------------------------------------------------------------- panels
def add_hazard(doc, h):
    lvl = h["level"]
    col = C.HAZARD_COLORS[lvl]
    bg = col["bg"].lstrip("#")
    fg = WHITE if col["fg"] == "#FFFFFF" else BLACK

    tbl = doc.add_table(rows=2, cols=1)
    tbl.autofit = False
    _set_borders(tbl, bg, sz=18)
    for r_ in tbl.rows:
        _cant_split(r_)

    # header band
    hc = tbl.rows[0].cells[0]
    _set_shade(hc, bg)
    _cell_margins(hc, top=50, bottom=50, left=120, right=120)
    hp = hc.paragraphs[0]
    hp.paragraph_format.space_after = Pt(0)
    # ANSI Z535: NOTICE carries no safety-alert symbol — signal word only.
    if lvl != "NOTICE":
        r = hp.add_run()
        r.add_picture(ic(h["icon"]), height=Inches(0.26))
        hp.add_run("  ")
    run(hp, lvl, size=14, bold=True, color=fg, spacing=40, caps=True)
    _keep_with_next(hp)

    # body
    bc = tbl.rows[1].cells[0]
    _set_shade(bc, "FFFFFF")
    _cell_margins(bc, top=70, bottom=80, left=120, right=120)
    tp = bc.paragraphs[0]
    tp.paragraph_format.space_after = Pt(2)
    run(tp, h["title"], size=10.5, bold=True, color=NAVY)
    bp = bc.add_paragraph()
    bp.paragraph_format.space_after = Pt(0)
    run(bp, h["text"], size=10)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ------------------------------------------------------------------- sections
def build():
    doc = Document()

    # base style
    normal = doc.styles["Normal"]
    normal.font.name = "Segoe UI"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = INK
    normal.paragraph_format.line_spacing = 1.04
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.space_before = Pt(0)

    sec = doc.sections[0]
    sec.page_height = Inches(11)
    sec.page_width = Inches(8.5)
    sec.top_margin = Inches(0.55)
    sec.bottom_margin = Inches(0.6)
    sec.left_margin = Inches(0.65)
    sec.right_margin = Inches(0.65)
    sec.header_distance = Inches(0.3)
    sec.footer_distance = Inches(0.3)

    # ---- running header ----
    hdr = sec.header
    htbl = hdr.add_table(rows=1, cols=2, width=Inches(7.2))
    htbl.autofit = True
    _no_borders(htbl)
    lc = htbl.rows[0].cells[0]
    lp = lc.paragraphs[0]
    lp.paragraph_format.space_after = Pt(0)
    lp.add_run().add_picture(ic("logo"), height=Inches(0.3))
    run(lp, "  ")
    run(lp, M["manufacturer"], size=15, bold=True, color=NAVY, spacing=20)
    sub = lc.add_paragraph()
    sub.paragraph_format.space_before = Pt(0)
    sub.paragraph_format.space_after = Pt(0)
    run(sub, M["subtitle"], size=7.5, color=GREY)
    rc = htbl.rows[0].cells[1]
    rp = rc.paragraphs[0]
    rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rp.paragraph_format.space_after = Pt(0)
    run(rp, "PRODUCT & SAFETY GUIDE", size=8, bold=True, color=WHITE, spacing=20)
    _set_shade(rc, "0E2A47")
    rc.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    # bottom rule under header
    hr = hdr.add_paragraph()
    hr.paragraph_format.space_before = Pt(2)
    hr.paragraph_format.space_after = Pt(0)
    _para_bottom_border(hr, "0E2A47", sz=18)

    # ---- running footer ----
    ftr = sec.footer
    ftbl = ftr.add_table(rows=1, cols=3, width=Inches(7.2))
    _no_borders(ftbl)
    f0 = ftbl.rows[0].cells[0].paragraphs[0]
    run(f0, f"{M['manufacturer']} {M['model']} — {M['product']}", size=7.5, color=GREY)
    f1 = ftbl.rows[0].cells[1].paragraphs[0]
    f1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(f1, f"{M['doc_no']} Rev {M['revision']} · {M['date']}", size=7.5, color=GREY)
    f2 = ftbl.rows[0].cells[2].paragraphs[0]
    f2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run(f2, "Page ", size=7.5, color=GREY)
    _field(f2, "PAGE")
    run(f2, " of ", size=7.5, color=GREY)
    _field(f2, "NUMPAGES")
    for r_ in (f0, f1, f2):
        _para_bottom_border  # noqa (top border instead below)
    # top rule above footer text
    # (done via the table cell top border)
    _set_borders(ftbl, LINE, sz=4, sides=("top",))

    # ===== PAGE 1 : overview + features =====
    tb = doc.add_paragraph()
    tb.paragraph_format.space_after = Pt(0)
    run(tb, M["model"], size=11, bold=True, color=TEAL, spacing=60)
    nm = doc.add_paragraph()
    nm.paragraph_format.space_before = Pt(0)
    nm.paragraph_format.space_after = Pt(0)
    run(nm, M["product"], size=30, bold=True, color=NAVY)
    sb = doc.add_paragraph()
    sb.paragraph_format.space_before = Pt(0)
    sb.paragraph_format.space_after = Pt(0)
    run(sb, M["subtitle"], size=12, color=INK)
    tg = doc.add_paragraph()
    tg.paragraph_format.space_before = Pt(1)
    run(tg, M["tagline"], size=10, italic=True, color=GREY)

    # hero + overview row
    hr = doc.add_table(rows=1, cols=2)
    _no_borders(hr)
    hr.columns[0].width = Inches(2.5)
    hr.columns[1].width = Inches(4.7)
    img_cell = hr.rows[0].cells[0]
    _set_shade(img_cell, "0E2A47")
    _cell_margins(img_cell, top=120, bottom=120, left=120, right=120)
    ip = img_cell.paragraphs[0]
    ip.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ip.add_run().add_picture(BOARD, height=Inches(2.7))
    ov_cell = hr.rows[0].cells[1]
    _cell_margins(ov_cell, left=180)
    op = ov_cell.paragraphs[0]
    op.paragraph_format.space_after = Pt(4)
    run(op, "Product Overview", size=12, bold=True, color=NAVY)
    op2 = ov_cell.add_paragraph()
    op2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run(op2, C.OVERVIEW, size=9.5)
    bp = ov_cell.add_paragraph()
    bp.paragraph_format.space_before = Pt(4)
    run(bp, "LTE Cat-1 + GPS   ·   LoRaWAN 915 MHz   ·   Bluetooth LE   ·   Universal I/O",
        size=8.5, bold=True, color=TEAL)

    heading(doc, "Key Features")
    ft = doc.add_table(rows=1, cols=3)
    _no_borders(ft)
    for idx, (title, items) in enumerate(C.FEATURES):
        cell = ft.rows[0].cells[idx]
        _cell_margins(cell, right=140)
        cp = cell.paragraphs[0]
        cp.paragraph_format.space_after = Pt(3)
        run(cp, title, size=10.5, bold=True, color=TEAL)
        for it_ in items:
            lp = cell.add_paragraph(style="List Bullet")
            lp.paragraph_format.space_after = Pt(2)
            run(lp, it_, size=9)

    doc.add_page_break()

    # ===== PAGE 2 : specifications + IO =====
    heading(doc, "Technical Specifications", space_before=0)
    st = doc.add_table(rows=len(C.SPECS), cols=2)
    _set_borders(st, LINE, sz=4)
    st.columns[0].width = Inches(2.6)
    st.columns[1].width = Inches(4.6)
    for i, (k, v) in enumerate(C.SPECS):
        kc, vc = st.rows[i].cells
        _set_shade(kc, SOFT)
        if i % 2:
            _set_shade(vc, "FBFCFD")
        _cell_margins(kc, top=24, bottom=24); _cell_margins(vc, top=24, bottom=24)
        kp = kc.paragraphs[0]; kp.paragraph_format.space_after = Pt(0)
        run(kp, k, size=9, bold=True, color=NAVY)
        vp = vc.paragraphs[0]; vp.paragraph_format.space_after = Pt(0)
        run(vp, v, size=9)

    heading(doc, "Field Interfaces at a Glance")
    it = doc.add_table(rows=len(C.IO_TABLE) + 1, cols=3)
    _set_borders(it, LINE, sz=4)
    it.columns[0].width = Inches(2.0)
    it.columns[1].width = Inches(0.7)
    it.columns[2].width = Inches(4.5)
    hdr_cells = it.rows[0].cells
    for j, htxt in enumerate(("Interface", "Qty", "Notes")):
        _set_shade(hdr_cells[j], "0E2A47")
        _cell_margins(hdr_cells[j], top=24, bottom=24)
        hp = hdr_cells[j].paragraphs[0]; hp.paragraph_format.space_after = Pt(0)
        if j == 1:
            hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(hp, htxt, size=9, bold=True, color=WHITE)
    for i, (fn, qty, note) in enumerate(C.IO_TABLE, start=1):
        cells = it.rows[i].cells
        if i % 2 == 0:
            for c in cells:
                _set_shade(c, "FBFCFD")
        for c in cells:
            _cell_margins(c, top=24, bottom=24)
        p0 = cells[0].paragraphs[0]; p0.paragraph_format.space_after = Pt(0)
        run(p0, fn, size=9, bold=True, color=NAVY)
        p1 = cells[1].paragraphs[0]; p1.paragraph_format.space_after = Pt(0)
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(p1, qty, size=9, bold=True)
        p2 = cells[2].paragraphs[0]; p2.paragraph_format.space_after = Pt(0)
        run(p2, note, size=9)

    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(6)
    run(note, "Note: the VX-0057 ships in Tank Monitor, Callbox, or RTU configurations. "
        "The interfaces exercised depend on the configuration ordered. Mechanical and "
        "environmental ratings are provided on the product label and in the deployment "
        "datasheet available from Viaanix.", size=9, italic=True, color=GREY)

    doc.add_page_break()

    # ===== PAGE 3+ : safety =====
    sh = doc.add_paragraph()
    sh.paragraph_format.space_after = Pt(8)
    run(sh, "Important Safety Information", size=18, bold=True, color=NAVY)
    _para_bottom_border(sh, "0E2A47", sz=22)

    # read-first banner
    rf = doc.add_table(rows=1, cols=2)
    _set_borders(rf, "F7C600", sz=18)
    rf.columns[0].width = Inches(0.7)
    rf.columns[1].width = Inches(6.5)
    icell, tcell = rf.rows[0].cells
    _set_shade(icell, "FFF8E1"); _set_shade(tcell, "FFF8E1")
    _cell_margins(icell, top=80, bottom=80); _cell_margins(tcell, top=80, bottom=80)
    icell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    ipp = icell.paragraphs[0]; ipp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ipp.paragraph_format.space_after = Pt(0)
    ipp.add_run().add_picture(ic("alert"), height=Inches(0.42))
    tpp = tcell.paragraphs[0]; tpp.paragraph_format.space_after = Pt(0)
    run(tpp, C.READ_FIRST, size=10, bold=True, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

    heading(doc, "Safety Symbols Used in This Document")
    sym = doc.add_table(rows=(len(C.SYMBOLS) + 1) // 2, cols=4)
    _no_borders(sym)
    sym.columns[0].width = Inches(0.55); sym.columns[1].width = Inches(3.05)
    sym.columns[2].width = Inches(0.55); sym.columns[3].width = Inches(3.05)
    for i, (name, term, meaning) in enumerate(C.SYMBOLS):
        r_ = i // 2
        base = (i % 2) * 2
        icl = sym.rows[r_].cells[base]
        tcl = sym.rows[r_].cells[base + 1]
        icl.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        _cell_margins(icl, top=40, bottom=40); _cell_margins(tcl, top=40, bottom=40)
        ipp = icl.paragraphs[0]; ipp.paragraph_format.space_after = Pt(0)
        ipp.add_run().add_picture(ic(name), height=Inches(0.34))
        t0 = tcl.paragraphs[0]; t0.paragraph_format.space_after = Pt(0)
        run(t0, term, size=9.5, bold=True, color=NAVY)
        t1 = tcl.add_paragraph(); t1.paragraph_format.space_after = Pt(0)
        run(t1, meaning, size=8.3, color=GREY)

    heading(doc, "Hazard Notices")
    for h in C.HAZARDS:
        add_hazard(doc, h)

    doc.add_page_break()

    # ===== handling =====
    hh = doc.add_paragraph()
    hh.paragraph_format.space_after = Pt(8)
    run(hh, "Handling, Installation & Disposal", size=18, bold=True, color=NAVY)
    _para_bottom_border(hh, "0E2A47", sz=22)

    for name, head, paras in C.HANDLING:
        t = doc.add_table(rows=1, cols=2)
        _no_borders(t)
        t.columns[0].width = Inches(0.55)
        t.columns[1].width = Inches(6.65)
        icl, tcl = t.rows[0].cells
        _cell_margins(icl, top=30); _cell_margins(tcl)
        ipp = icl.paragraphs[0]; ipp.paragraph_format.space_after = Pt(0)
        ipp.add_run().add_picture(ic(name), height=Inches(0.32))
        hp = tcl.paragraphs[0]; hp.paragraph_format.space_after = Pt(1)
        run(hp, head, size=10.5, bold=True, color=TEAL)
        for p in paras:
            pp = tcl.add_paragraph()
            pp.paragraph_format.space_after = Pt(0)
            run(pp, p, size=9.5)
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)
        sp.paragraph_format.space_before = Pt(0)
        sp.add_run().font.size = Pt(4)

    # compliance box
    cb = doc.add_table(rows=1, cols=2)
    _set_borders(cb, LINE, sz=4)
    cb.columns[0].width = Inches(0.6); cb.columns[1].width = Inches(6.6)
    icl, tcl = cb.rows[0].cells
    _set_shade(icl, SOFT); _set_shade(tcl, SOFT)
    _cell_margins(icl, top=90, bottom=90); _cell_margins(tcl, top=90, bottom=90)
    icl.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    ipp = icl.paragraphs[0]; ipp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ipp.paragraph_format.space_after = Pt(0)
    ipp.add_run().add_picture(ic("book"), height=Inches(0.34))
    c0 = tcl.paragraphs[0]; c0.paragraph_format.space_after = Pt(2)
    run(c0, "Compliance & Certification", size=10.5, bold=True, color=NAVY)
    c1 = tcl.add_paragraph(); c1.paragraph_format.space_after = Pt(0)
    run(c1, C.COMPLIANCE, size=9.2)

    end = doc.add_paragraph()
    end.alignment = WD_ALIGN_PARAGRAPH.CENTER
    end.paragraph_format.space_before = Pt(12)
    run(end, M["copyright"], size=8, color=GREY)

    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    p = build()
    print("DOCX:", p, "exists" if os.path.exists(p) else "MISSING",
          f"({os.path.getsize(p)} bytes)")
